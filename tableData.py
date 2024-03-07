import subprocess
import sys

from docx import Document
def extract_tables_with_heading(doc):
    tables = []

    for table in doc.tables:
        # Check if the heading "field descriptions" is present in the first row
        first_row_text = table.rows[0].cells[0].text.strip().lower()
        sub_ie = ""
        if "field descriptions" in first_row_text:
            sub_ie = first_row_text.split('field descriptions')[0].strip()
            # Extract and add the table to the list
            rows = []
            for row in table.rows[1:]:
                rows.append("<start>")
                rows.append(sub_ie)
                cell_texts = [cell.text.strip() for cell in row.cells]
                rows.append("\t".join(cell_texts))
                rows.append("<endl>")
                rows.append("\n")
            tables.append("\n".join(rows))

    return tables

def save_tables_to_doc(tables, output_path):
    doc = Document()

    for i, table in enumerate(tables, start=1):
        doc.add_paragraph(table)
        # doc.add_paragraph("")  # Add an extra line break

    doc.save(output_path)


def main(docx_filename, version):
    # Debugging
    input_document_path = docx_filename
    output_document_path = "extracted_tables_with_heading.docx"

    # Open the input document
    input_doc = Document(input_document_path)

    # Extract tables with the heading "field descriptions"
    tables = extract_tables_with_heading(input_doc)

    # Save tables to a new document with extra line breaks
    save_tables_to_doc(tables, output_document_path)

    subprocess.run(["python", "storingTableData.py", docx_filename, version])


if __name__ == "__main__":
    # print("Hello world")
    if len(sys.argv) != 3:
        print("Usage: python tableData <docx_filename>")
        # sys.exit(1)
    docx_filename = sys.argv[1]
    version = sys.argv[2]
    main(docx_filename, version)
