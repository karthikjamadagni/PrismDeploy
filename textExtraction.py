import sqlite3
import subprocess
import sys

from docx import Document
import re


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def save_text_to_docx(text, output_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)


def extract_ie_structures(text):
    # Define a pattern for identifying ASN.1 structures
    asn1_pattern = re.compile(r'(-- ASN1START\n-- TAG-(.*?)-START\n\n(.*?)\n\n-- TAG-\2-STOP\n-- ASN1STOP)', re.DOTALL)

    # Extract IE name and ASN.1 structures from the text
    ie_matches = asn1_pattern.findall(text)

    return ie_matches


def save_ie_matches_to_docx(ie_matches, output_path):
    doc = Document()

    for match in ie_matches:
        ie_structure = match[0]  # Full matched substring
        doc.add_paragraph(ie_structure + '\n')

    doc.save(output_path)


def extract_and_store_ie_info(text, output_path, database_path, version):
    output_doc = Document()

    # Define a pattern for identifying ASN.1 structures
    asn1_pattern = re.compile(r'-- ASN1START\n-- TAG-(.*?)-START\s*(.*?)\s*-- TAG-\1-STOP\n-- ASN1STOP', re.DOTALL)

    # Extract IE name and ASN.1 structures from the text
    ie_matches = asn1_pattern.findall(text)

    # Create a SQLite database and a table for IE information
    conn = sqlite3.connect(database_path)
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS ie_info (
                                id INTEGER PRIMARY KEY AUTOINCREMENT,
                                ie_name TEXT,
                                sub_ie TEXT,
                                field TEXT,
                                type TEXT,
                                description TEXT,
                                status BOOLEAN DEFAULT 0,
                                ver TEXT
                             )''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS ie_structures (
                                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                                    ie_name TEXT,
                                    structure TEXT,
                                    ver TEXT
                                 )''')
    conn.commit()

    for ie_match in ie_matches:
        ie_name = ie_match[0].strip()
        ie_structure = ie_match[1].strip()

        cursor.execute("INSERT INTO ie_structures (ie_name, structure, ver) VALUES (?, ?, ?)", (ie_name, ie_structure, version))
        conn.commit()
        # Add IE name and fields to the output document
        output_doc.add_paragraph(f"Name of the IE: ({ie_name})")
        sub_ie = ""
        # Extract information from the lines after IE name
        info_lines = ie_match[1].split('\n')
        for info_line in info_lines:
            info_line = info_line.strip()
            # Ignore lines containing only '{' or '}' or empty lines
            if not (info_line.strip() == '{' or info_line.strip() == '}' or not info_line.strip()):
                if "::=" in info_line:
                    sub_ie = info_line.split('::=')[0].strip()
                    if "SEQUENCE" in info_line:
                        after_ = info_line.split('SEQUENCE', 1)[1].strip()
                        if not after_ == "{":
                            output_doc.add_paragraph(f"Name of the Sub_IE: ({sub_ie})")
                            temp = sub_ie + " SEQUENCE " + after_
                            typeandcondition = " SEQUENCE " + after_
                            output_doc.add_paragraph(temp)
                            cursor.execute("INSERT INTO ie_info (ie_name, sub_ie, field, type, ver) VALUES (?, ?, ?, ?, ?)",
                                           (ie_name, sub_ie, sub_ie, typeandcondition, version))
                    elif "ENUMERATED" in info_line:
                        after_ = info_line.split('ENUMERATED', 1)[1].strip()
                        if not after_ == "{":
                            output_doc.add_paragraph(f"Name of the Sub_IE: ({sub_ie})")
                            temp = sub_ie + " ENUMERATED " + after_
                            typeandcondition = " SEQUENCE " + after_
                            output_doc.add_paragraph(temp)
                            cursor.execute("INSERT INTO ie_info (ie_name, sub_ie, field, type, ver) VALUES (?, ?, ?, ?, ?)",
                                           (ie_name, sub_ie, sub_ie, typeandcondition, version))
                else:
                    output_doc.add_paragraph(f"Name of the Sub_IE: ({sub_ie})")
                    output_doc.add_paragraph(info_line)
                    info_line = info_line.strip()
                    parts = info_line.split(maxsplit=1)
                    field = parts[0]
                    typeandcondition = parts[1] if len(parts) > 1 else ''
                    cursor.execute("INSERT INTO ie_info (ie_name, sub_ie, field, type, ver) VALUES (?, ?, ?, ?, ?)",
                                   (ie_name, sub_ie, field, typeandcondition, version))

    output_doc.save(output_path)
    conn.commit()
    conn.close()


def main(docx_filename, version):
    import time
    start_time = time.time()
    # Debugging
    input_document_path = docx_filename
    output_document_path = "test.docx"
    database_path = "ie_info.db"

    # Extract text from the full document
    document_text = extract_text_from_docx(input_document_path)

    # Extract and store IE information in the output document and the database
    extract_and_store_ie_info(document_text, output_document_path, database_path, version)

    subprocess.run(["python", "tableData.py", docx_filename, version])

    end_time = time.time()
    elapsed_time = end_time - start_time

    print(f"The program took {elapsed_time:.2f} seconds to run.")


if __name__ == "__main__":
    # print("Hello world")
    # if len(sys.argv) != 3:
    #     print("Usage: python textExtraction.py <docx_filename>")
        # sys.exit(1)
    docx_filename = sys.argv[1]
    version = sys.argv[2]
    print(docx_filename)
    print(version)
    # docx_filename = "3gpp_specs\\38331-h00.docx"
    # version = "17.1.0"
    main(docx_filename, version)
    # print("....", docx_filename, "....")
    # print(version)
